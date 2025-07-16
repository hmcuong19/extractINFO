import streamlit as st
import google.generativeai as genai
import io
import docx
import fitz  # PyMuPDF

# --- Cấu hình và Thiết lập ---

# Thiết lập tiêu đề và icon cho trang, sử dụng layout rộng để có 2 cột
st.set_page_config(page_title="Trích xuất Thông tin Thông minh", page_icon="✨", layout="wide")

# Lấy API key từ secrets của Streamlit để bảo mật
# Hướng dẫn: https://docs.streamlit.io/deploy/streamlit-community-cloud/deploy-your-app/secrets-management
try:
    # Cố gắng lấy key từ Streamlit's secrets management
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=GOOGLE_API_KEY)
except (KeyError, FileNotFoundError):
    # Nếu không tìm thấy trong secrets (khi chạy local), yêu cầu người dùng nhập thủ công
    st.warning("Không tìm thấy Google API Key trong Streamlit secrets. Vui lòng nhập thủ công để chạy ứng dụng.")
    GOOGLE_API_KEY = st.text_input("Nhập Google API Key của bạn:", type="password")
    if GOOGLE_API_KEY:
        genai.configure(api_key=GOOGLE_API_KEY)
    else:
        # Dừng ứng dụng nếu không có key để tránh lỗi
        st.info("Vui lòng cung cấp API key để bắt đầu.")
        st.stop()

# --- Các hàm xử lý ---

def get_gemini_response(input_text, prompt):
    """
    Hàm gọi Gemini API để lấy phản hồi dựa trên văn bản và prompt.
    Sử dụng model 'gemini-1.5-flash' là phiên bản mới và hiệu quả.
    """
    model = genai.GenerativeModel('gemini-1.5-flash')
    try:
        response = model.generate_content([input_text, prompt])
        return response.text
    except Exception as e:
        # Bắt lỗi và trả về thông báo thân thiện
        return f"Đã xảy ra lỗi khi gọi API Gemini: {e}"

def _get_text_from_docx(docx_bytes):
    """
    Hàm nội bộ để trích xuất văn bản thô từ file .docx (dạng bytes).
    CẬP NHẬT: Nâng cấp để đọc cả nội dung trong bảng (tables), không chỉ các đoạn văn thông thường.
    """
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        full_text = []
        # Lặp qua các đoạn văn bản
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Lặp qua các bảng
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Lỗi đọc file .docx: {e}")
        return None

def convert_docx_to_pdf(docx_bytes):
    """
    Chuyển đổi file .docx (dưới dạng bytes) sang file .pdf (dưới dạng bytes).
    Lưu ý: Đây là chuyển đổi dựa trên văn bản, bố cục gốc của file .docx sẽ không được giữ lại.
    """
    try:
        # Bước 1: Trích xuất văn bản từ file .docx
        text = _get_text_from_docx(docx_bytes)
        if text is None:
            return None

        # Bước 2: Tạo một file PDF mới trong bộ nhớ
        pdf_doc = fitz.open()
        # Thêm một trang với kích thước A4
        page = pdf_doc.new_page(width=595, height=842)

        # Bước 3: Chèn văn bản đã trích xuất vào trang PDF
        # insert_textbox sẽ tự động xử lý việc xuống dòng và ngắt trang cơ bản
        page.insert_textbox(fitz.Rect(50, 50, 545, 792), text, fontsize=11, fontname="helv", align=fitz.TEXT_ALIGN_LEFT)

        # Bước 4: Lưu file PDF ra dưới dạng bytes
        pdf_bytes = pdf_doc.tobytes()
        pdf_doc.close()
        return pdf_bytes
    except Exception as e:
        st.error(f"Lỗi khi chuyển đổi DOCX sang PDF: {e}")
        return None


def extract_text_from_pdf(file_bytes):
    """
    Hàm trích xuất toàn bộ văn bản từ file .pdf (dạng bytes).
    """
    try:
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        full_text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            full_text += page.get_text()
        pdf_document.close()
        return full_text
    except Exception as e:
        st.error(f"Lỗi đọc file .pdf: {e}")
        return None

# --- Giao diện ứng dụng Streamlit ---

st.title("✨ Trích xuất Thông tin từ Tài liệu với Gemini Pro")
st.markdown("Tải lên tệp `.docx` hoặc `.pdf`. Tệp `.docx` sẽ được tự động chuyển sang `.pdf` trước khi xử lý.")

# Tạo hai cột với tỉ lệ chiều rộng 2:3
col1, col2 = st.columns([2, 3])

with col1:
    st.header("1. Tải lên & Tùy chỉnh")

    # Widget tải file
    uploaded_file = st.file_uploader("Chọn một tệp (.docx hoặc .pdf)", type=['docx', 'pdf'])

    # Prompt mặc định
    prompt_default = """Bạn là một trợ lý AI chuyên nghiệp trong việc trích xuất thông tin.
Dựa vào nội dung văn bản được cung cấp, hãy tách và liệt kê các thông tin sau:
Từ nội dung đề cương học phần dưới đây, hãy trích xuất và trình bày rõ ràng các mục sau:
Tên học phần
Mã học phần (nếu có)
Số tín chỉ
Điều kiện tiên quyết (nếu có)
Mục tiêu học phần
Chuẩn đầu ra của học phần (CLO)
Nội dung học phần tóm tắt
Tài liệu tham khảo (ghi rõ tên, tác giả, năm, NXB nếu có)

Trình bày câu trả lời theo định dạng rõ ràng như sau:
Tên học phần: ...
Mã học phần: ...
Số tín chỉ: ...
Điều kiện tiên quyết: ...
Mục tiêu học phần:
- ...
- ...
Chuẩn đầu ra:
- CLO1: ...
- CLO2: ...
...
Tóm tắt nội dung học phần:
- Tuần 1: ...
- Tuần 2: ...
...
Tài liệu tham khảo:
- ...
- ...

Nếu không tìm thấy thông tin nào, hãy ghi là "Không tìm thấy".
"""
    prompt_user = st.text_area("Chỉnh sửa prompt (câu lệnh yêu cầu):", value=prompt_default, height=350)

    # Nút xử lý
    submit_button = st.button("🚀 Bắt đầu trích xuất")

with col2:
    st.header("2. Kết quả trích xuất")

    # Vùng chứa kết quả
    result_container = st.container()
    result_container.info("Kết quả sẽ được hiển thị ở đây sau khi bạn nhấn nút 'Bắt đầu trích xuất'.")

    # Xử lý logic khi người dùng nhấn nút
    if submit_button:
        if uploaded_file is not None and prompt_user:
            with st.spinner("Đang xử lý file... Vui lòng chờ! 🤖"):
                file_bytes = uploaded_file.getvalue()
                file_extension = uploaded_file.name.split('.')[-1].lower()
                
                pdf_for_processing = None
                
                if file_extension == "docx":
                    st.info("Phát hiện file .docx. Đang tiến hành chuyển đổi sang .pdf...")
                    pdf_for_processing = convert_docx_to_pdf(file_bytes)
                    if pdf_for_processing:
                        st.success("Chuyển đổi thành công!")
                    else:
                        st.error("Lỗi trong quá trình chuyển đổi .docx sang .pdf.")
                        st.stop()
                elif file_extension == "pdf":
                    pdf_for_processing = file_bytes

                # Trích xuất văn bản từ file pdf (gốc hoặc đã chuyển đổi)
                st.info("Đang trích xuất văn bản...")
                raw_text = extract_text_from_pdf(pdf_for_processing)
                
                if raw_text and raw_text.strip(): # Dùng strip() để đảm bảo chuỗi không chỉ chứa khoảng trắng
                    st.info("Văn bản đã được trích xuất. Đang gửi yêu cầu đến Gemini...")
                    response = get_gemini_response(raw_text, prompt_user)
                    result_container.text_area("Thông tin đã trích xuất:", value=response, height=550)
                elif raw_text is not None: # Trường hợp raw_text là chuỗi rỗng ""
                    result_container.warning("Không tìm thấy nội dung văn bản nào trong file. File có thể chỉ chứa hình ảnh hoặc không có văn bản để trích xuất.")
                else: # Trường hợp raw_text là None, tức là có lỗi xảy ra trong hàm extract_text_from_pdf
                    result_container.error("Không thể đọc được nội dung từ file do có lỗi xảy ra trong quá trình xử lý. Vui lòng thử lại với file khác.")
        elif not uploaded_file:
            st.warning("Vui lòng tải lên một file để tiếp tục.")
        else:
            st.warning("Prompt không được để trống.")
